import pandas as pd
from scipy.stats import spearmanr
from docx import Document
from docx.shared import Inches

# Load the dataset
file_path = r"D:\RESEARCH\Water Quality\Drinking.csv"
data = pd.read_csv(file_path)

# Calculate Spearman correlation matrix
corr_matrix, _ = spearmanr(data)

# Convert the correlation matrix to a DataFrame for better handling
corr_df = pd.DataFrame(corr_matrix, columns=data.columns, index=data.columns)

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading('Spearman Correlation Matrix', level=1)

# Create a table in the Word document
table = doc.add_table(rows=corr_df.shape[0] + 1, cols=corr_df.shape[1] + 1)

# Add headers to the table
table.cell(0, 0).text = 'Variable'
for i, col in enumerate(corr_df.columns):
    table.cell(0, i + 1).text = col

# Fill the table with correlation values
for i, row in enumerate(corr_df.index):
    table.cell(i + 1, 0).text = row
    for j, col in enumerate(corr_df.columns):
        table.cell(i + 1, j + 1).text = f"{corr_df.iloc[i, j]:.2f}"

# Save the Word document
output_path = r"D:\RESEARCH\Water Quality\Spearman_Correlation_Matrix.docx"
doc.save(output_path)

print(f"Spearman Correlation Matrix saved to {output_path}")